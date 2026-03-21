"""Tests for service/engine.py — no API tokens burned.

Covers: dataset loading, resume logic, result rows, thinking extraction,
score_rubric, judge pipeline (mocked), result-worker queue lifecycle,
on_case_complete callback, write_meta atomicity, storage snapshot sync,
and full end-to-end async eval with FakeLLMClient.
"""
from __future__ import annotations

import asyncio
import json
import os
import textwrap
import time
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from runner import CaseResult, EvalCase

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_task(
    tid: str = "t1",
    task: str = "solve it",
    rubric: dict | None = None,
    **extra: Any,
) -> dict[str, Any]:
    d: dict[str, Any] = {"id": tid, "task": task}
    if rubric is not None:
        d["rubric"] = rubric
    d.update(extra)
    return d


def _make_result(
    case_id: str = "t1",
    status: str = "ok",
    output_text: str = "answer",
) -> CaseResult:
    return CaseResult(
        case_id=case_id,
        status=status,
        output_text=output_text,
        error=None if status == "ok" else "some error",
        model_wait_s=1.0,
        tool_cpu_s=0.1,
        total_s=1.1,
        model_calls=2,
        tool_calls=1,
        peak_in_flight_evals=4,
        peak_in_flight_tools=2,
        started_at_s=0.0,
        finished_at_s=1.1,
    )


# ---------------------------------------------------------------------------
# load_dataset
# ---------------------------------------------------------------------------

class TestLoadDataset:
    def test_loads_jsonl(self, tmp_path: Path) -> None:
        from service.engine import load_dataset

        p = tmp_path / "ds.jsonl"
        p.write_text(
            json.dumps({"id": "a", "task": "x"}) + "\n"
            + json.dumps({"id": "b", "task": "y"}) + "\n"
        )
        rows = load_dataset(p)
        assert len(rows) == 2
        assert rows[0]["id"] == "a"
        assert rows[1]["id"] == "b"

    def test_loads_json_array(self, tmp_path: Path) -> None:
        from service.engine import load_dataset

        p = tmp_path / "ds.json"
        p.write_text(json.dumps([{"id": "c"}, {"id": "d"}]))
        rows = load_dataset(p)
        assert len(rows) == 2

    def test_json_non_list_raises(self, tmp_path: Path) -> None:
        from service.engine import load_dataset

        p = tmp_path / "ds.json"
        p.write_text(json.dumps({"not": "a list"}))
        with pytest.raises(ValueError, match="list"):
            load_dataset(p)

    def test_skips_blank_lines(self, tmp_path: Path) -> None:
        from service.engine import load_dataset

        p = tmp_path / "ds.jsonl"
        p.write_text(json.dumps({"id": "a"}) + "\n\n" + json.dumps({"id": "b"}) + "\n")
        assert len(load_dataset(p)) == 2

    def test_bad_json_raises(self, tmp_path: Path) -> None:
        from service.engine import load_dataset

        p = tmp_path / "ds.jsonl"
        p.write_text("not json\n")
        with pytest.raises(ValueError, match="Invalid JSON"):
            load_dataset(p)


# ---------------------------------------------------------------------------
# _extract_completed  (resume logic)
# ---------------------------------------------------------------------------

class TestExtractCompleted:
    def test_empty_file(self, tmp_path: Path) -> None:
        from service.engine import _extract_completed

        ids, ok, fail = _extract_completed(tmp_path / "nope.jsonl")
        assert ids == set()
        assert ok == 0 and fail == 0

    def test_counts_ok_and_failed(self, tmp_path: Path) -> None:
        from service.engine import _extract_completed

        p = tmp_path / "out.jsonl"
        rows = [
            {"id": "a", "status": "ok"},
            {"id": "b", "status": "error"},
            {"id": "c", "eval": {"score": 0.5}},  # no status → counted as ok
        ]
        p.write_text("\n".join(json.dumps(r) for r in rows) + "\n")
        ids, ok, fail = _extract_completed(p)
        assert ids == {"a", "b", "c"}
        assert ok == 2
        assert fail == 1

    def test_skips_bad_json(self, tmp_path: Path) -> None:
        from service.engine import _extract_completed

        p = tmp_path / "out.jsonl"
        p.write_text('bad\n{"id": "x", "status": "ok"}\n')
        ids, ok, fail = _extract_completed(p)
        assert ids == {"x"}
        assert ok == 1

    def test_uses_case_id_fallback(self, tmp_path: Path) -> None:
        from service.engine import _extract_completed

        p = tmp_path / "out.jsonl"
        p.write_text(json.dumps({"case_id": "z", "status": "ok"}) + "\n")
        ids, _, _ = _extract_completed(p)
        assert "z" in ids


# ---------------------------------------------------------------------------
# _write_meta
# ---------------------------------------------------------------------------

class TestWriteMeta:
    def test_creates_new(self, tmp_path: Path) -> None:
        from service.engine import _write_meta

        p = tmp_path / "m.meta.json"
        _write_meta(p, status="running", count=5)
        data = json.loads(p.read_text())
        assert data["status"] == "running"
        assert data["count"] == 5

    def test_merges_existing(self, tmp_path: Path) -> None:
        from service.engine import _write_meta

        p = tmp_path / "m.meta.json"
        _write_meta(p, status="running", ok=0)
        _write_meta(p, status="completed", ok=10)
        data = json.loads(p.read_text())
        assert data["status"] == "completed"
        assert data["ok"] == 10

    def test_atomic_via_rename(self, tmp_path: Path) -> None:
        from service.engine import _write_meta

        p = tmp_path / "m.meta.json"
        _write_meta(p, x=1)
        # .tmp should not persist
        assert not (tmp_path / "m.tmp").exists()


# ---------------------------------------------------------------------------
# _result_row
# ---------------------------------------------------------------------------

class TestResultRow:
    def test_basic_fields(self) -> None:
        from service.engine import _result_row

        task = _make_task("t1", "solve it", source="bench", category="math")
        result = _make_result("t1")
        row = _result_row(task=task, result=result, llm_id="test-llm")

        assert row["id"] == "t1"
        assert row["case_id"] == "t1"
        assert row["llm_answer"] == "answer"
        assert row["status"] == "ok"
        assert row["runner"] == "async"
        assert row["llm_id"] == "test-llm"
        assert row["source"] == "bench"
        assert row["category"] == "math"
        assert row["metrics"]["model_wait_s"] == 1.0
        assert row["metrics"]["tool_calls"] == 1

    def test_preserves_task_keys(self) -> None:
        from service.engine import _result_row

        task = _make_task("t1", "solve", custom_field="hello")
        row = _result_row(task=task, result=_make_result(), llm_id="x")
        assert row["custom_field"] == "hello"

    def test_no_metadata_when_none(self) -> None:
        from service.engine import _result_row

        row = _result_row(
            task=_make_task(), result=_make_result(), llm_id="x", llm_metadata=None
        )
        assert "llm_metadata" not in row

    def test_includes_metadata(self) -> None:
        from service.engine import _result_row

        meta = {"usage": {"total_tokens": 100}, "thinking": "deep thought"}
        row = _result_row(
            task=_make_task(), result=_make_result(), llm_id="x", llm_metadata=meta
        )
        assert row["llm_metadata"] == meta
        assert row["thinking"] == "deep thought"


# ---------------------------------------------------------------------------
# _extract_thinking
# ---------------------------------------------------------------------------

class TestExtractThinking:
    def test_none_metadata(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking(None) is None

    def test_not_a_dict(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking("string") is None  # type: ignore[arg-type]

    def test_direct_thinking_string(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking({"thinking": "hmm"}) == "hmm"

    def test_direct_reasoning_list(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking({"reasoning": ["step1", "step2"]}) == ["step1", "step2"]

    def test_empty_thinking_ignored(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking({"thinking": "  "}) is None

    def test_turns_extraction(self) -> None:
        from service.engine import _extract_thinking

        meta = {
            "turns": [
                {"turn": 0, "thinking": "first"},
                {"turn": 1, "content": "no thinking"},
                {"turn": 2, "thinking": "third"},
            ]
        }
        result = _extract_thinking(meta)
        assert len(result) == 2
        assert result[0] == {"turn": 0, "thinking": "first"}
        assert result[1] == {"turn": 2, "thinking": "third"}

    def test_turns_without_turn_index(self) -> None:
        from service.engine import _extract_thinking

        meta = {"turns": [{"thinking": "solo"}]}
        result = _extract_thinking(meta)
        assert result == [{"thinking": "solo"}]

    def test_empty_turns_returns_none(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking({"turns": []}) is None

    def test_turns_no_thinking_returns_none(self) -> None:
        from service.engine import _extract_thinking
        assert _extract_thinking({"turns": [{"content": "nothing"}]}) is None


# ---------------------------------------------------------------------------
# score_rubric (from eval/core.py — pure math, no API)
# ---------------------------------------------------------------------------

class TestScoreRubric:
    def test_all_pass(self) -> None:
        from eval.core import score_rubric

        score = score_rubric([True, True], [True, True], [True])
        assert score == 1.0

    def test_mandatory_fail_zero(self) -> None:
        from eval.core import score_rubric

        assert score_rubric([True, False], [True], [True]) == 0.0

    def test_empty_mandatory_gives_base(self) -> None:
        from eval.core import score_rubric

        # all([]) is True → base 0.40
        assert score_rubric([], [], []) == 0.40

    def test_partial_good_to_have(self) -> None:
        from eval.core import score_rubric

        # base 0.40 + 0.35 * (1/2) = 0.575
        score = score_rubric([True], [True, False], [])
        assert score == pytest.approx(0.575)

    def test_partial_ideal(self) -> None:
        from eval.core import score_rubric

        # base 0.40 + 0.25 * (1/2) = 0.525
        score = score_rubric([True], [], [True, False])
        assert score == pytest.approx(0.525)


# ---------------------------------------------------------------------------
# _to_case
# ---------------------------------------------------------------------------

class TestToCase:
    def test_basic_conversion(self) -> None:
        from service.engine import AsyncRunConfig, _to_case

        task = _make_task("t1", "do thing", source="bench", category="math")
        case = _to_case(task, 0, AsyncRunConfig())
        assert case.case_id == "t1"
        assert case.prompt == "do thing"
        assert case.tool_mode == "sleep"
        assert case.max_steps == 3

    def test_task_overrides_config(self) -> None:
        from service.engine import AsyncRunConfig, _to_case

        task = _make_task(tool_mode="cpu", tool_payload=5, max_steps=10)
        case = _to_case(task, 0, AsyncRunConfig(case_tool_mode="sleep", case_max_steps=3))
        assert case.tool_mode == "cpu"
        assert case.tool_payload == 5
        assert case.max_steps == 10

    def test_fallback_id(self) -> None:
        from service.engine import AsyncRunConfig, _to_case

        task = {"task": "no id"}
        case = _to_case(task, 7, AsyncRunConfig())
        assert case.case_id == "case-0007"


# ---------------------------------------------------------------------------
# _task_id / _task_prompt
# ---------------------------------------------------------------------------

class TestTaskHelpers:
    def test_task_id_from_id(self) -> None:
        from service.engine import _task_id
        assert _task_id({"id": "abc"}, 0) == "abc"

    def test_task_id_fallback(self) -> None:
        from service.engine import _task_id
        assert _task_id({}, 5) == "case-0005"

    def test_task_prompt_from_task(self) -> None:
        from service.engine import _task_prompt
        assert _task_prompt({"task": "do this"}) == "do this"

    def test_task_prompt_from_prompt(self) -> None:
        from service.engine import _task_prompt
        assert _task_prompt({"prompt": "do that"}) == "do that"

    def test_task_prompt_empty(self) -> None:
        from service.engine import _task_prompt
        assert _task_prompt({}) == ""


# ---------------------------------------------------------------------------
# _wandb_usage_tokens
# ---------------------------------------------------------------------------

class TestWandbUsageTokens:
    def test_none(self) -> None:
        from service.engine import _wandb_usage_tokens
        assert _wandb_usage_tokens(None) is None

    def test_total_tokens(self) -> None:
        from service.engine import _wandb_usage_tokens
        assert _wandb_usage_tokens({"usage": {"total_tokens": 500}}) == 500

    def test_input_output_tokens(self) -> None:
        from service.engine import _wandb_usage_tokens

        meta = {"usage": {"input_tokens": 100, "output_tokens": 200}}
        assert _wandb_usage_tokens(meta) == 300

    def test_prompt_completion_tokens(self) -> None:
        from service.engine import _wandb_usage_tokens

        meta = {"usage": {"prompt_tokens": 50, "completion_tokens": 75}}
        assert _wandb_usage_tokens(meta) == 125

    def test_no_usage_key(self) -> None:
        from service.engine import _wandb_usage_tokens
        assert _wandb_usage_tokens({"other": 1}) is None

    def test_zero_tokens_returns_none(self) -> None:
        from service.engine import _wandb_usage_tokens

        meta = {"usage": {"input_tokens": 0, "output_tokens": 0}}
        assert _wandb_usage_tokens(meta) is None


# ---------------------------------------------------------------------------
# _build_judged_answer (mock output files)
# ---------------------------------------------------------------------------

class TestBuildJudgedAnswer:
    def test_no_output_dir(self) -> None:
        from service.engine import _build_judged_answer

        task = _make_task()
        assert _build_judged_answer(task, "my answer") == "my answer"

    def test_with_output_dir(self, tmp_path: Path) -> None:
        from service.engine import _build_judged_answer

        out_dir = tmp_path / "out"
        out_dir.mkdir()
        (out_dir / "result.txt").write_text("hello world")

        task = _make_task(output_dir=str(out_dir.relative_to(tmp_path)))
        with patch("eval.core._resolve_output_dir", return_value=out_dir):
            result = _build_judged_answer(task, "answer")
        assert "answer" in result
        assert "hello world" in result
        assert "ARTIFACT CONTEXT" in result


class TestOutputArtifactHandling:
    def test_collect_output_files_recurses_under_artifact_root(self, tmp_path: Path) -> None:
        from eval.core import _collect_output_files

        out_dir = tmp_path / "artifact"
        nested = out_dir / "nested"
        nested.mkdir(parents=True)
        (nested / "report.txt").write_text("nested hello")

        collected = _collect_output_files(out_dir)

        assert "nested/report.txt" in collected["inline"]
        assert "nested hello" in collected["inline"]

    def test_build_judge_artifact_bundle_collects_manifest_and_preloads(self, tmp_path: Path) -> None:
        from eval.core import build_judge_artifact_bundle

        out_dir = tmp_path / "artifact"
        nested = out_dir / "nested"
        nested.mkdir(parents=True)
        (nested / "report.txt").write_text("nested hello")

        with patch("eval.core._resolve_output_dir", return_value=out_dir):
            bundle = build_judge_artifact_bundle("artifact")

        assert bundle["artifact_root"] == str(out_dir.resolve())
        manifest_paths = [entry["relative_path"] for entry in bundle["manifest"]]
        assert manifest_paths == ["nested/report.txt"]
        assert "nested/report.txt" in bundle["prompt_context"]
        assert "nested hello" in bundle["prompt_context"]
        assert bundle["repl_seed"]["artifact_root"] == str(out_dir.resolve())
        assert bundle["repl_seed"]["artifacts_text"]["nested/report.txt"] == "nested hello"
        assert bundle["repl_seed"]["nested_report_text"] == "nested hello"

    def test_prepare_task_output_dir_clears_nested_children(self, tmp_path: Path) -> None:
        from service.engine import _prepare_task_output_dir

        out_dir = tmp_path / "artifact" / "case-1"
        nested = out_dir / "old"
        nested.mkdir(parents=True)
        (nested / "stale.txt").write_text("old data")

        task = _make_task(tid="case-1", output_dir="artifact")
        with patch("eval.core._resolve_output_dir", return_value=out_dir):
            prepared = _prepare_task_output_dir(task)

        assert prepared == str(out_dir)
        assert out_dir.exists()
        assert list(out_dir.iterdir()) == []


# ---------------------------------------------------------------------------
# _judge_task_with_rubric (mock the Gemini call)
# ---------------------------------------------------------------------------

class TestJudgeTaskWithRubric:
    def test_no_rubric_returns_none(self) -> None:
        from service.engine import _judge_task_with_rubric

        assert _judge_task_with_rubric({"task": "x"}, "ans", criterion_workers=1) is None

    def test_empty_rubric_returns_none(self) -> None:
        from service.engine import _judge_task_with_rubric

        task = _make_task(rubric={"mandatory": [], "good_to_have": [], "ideal": []})
        assert _judge_task_with_rubric(task, "ans", criterion_workers=1) is None

    def test_rubric_with_mocked_judge(self) -> None:
        from service.engine import _judge_task_with_rubric

        task = _make_task(rubric={
            "mandatory": ["criterion1"],
            "good_to_have": ["criterion2"],
            "ideal": [],
        })
        mock_judge = {"mandatory": [True], "good_to_have": [True], "ideal": []}
        with patch("eval.core.judge_rubric", return_value=mock_judge):
            with patch("eval.core.score_rubric", return_value=0.75):
                result = _judge_task_with_rubric(task, "answer", criterion_workers=2)

        assert result is not None
        assert result["score"] == 0.75
        assert result["mandatory"] == [True]

    def test_rubric_only_mandatory(self) -> None:
        from service.engine import _judge_task_with_rubric

        task = _make_task(rubric={"mandatory": ["m1", "m2"]})
        mock_judge = {"mandatory": [True, False], "good_to_have": [], "ideal": []}
        with patch("eval.core.judge_rubric", return_value=mock_judge):
            with patch("eval.core.score_rubric", return_value=0.0):
                result = _judge_task_with_rubric(task, "ans", criterion_workers=1)
        assert result["score"] == 0.0

    def test_judge_receives_artifact_dir_for_repl(self, tmp_path: Path) -> None:
        from service.engine import _judge_task_with_rubric

        out_dir = tmp_path / "artifact" / "case-1"
        out_dir.mkdir(parents=True)
        (out_dir / "summary.txt").write_text("hello judge")
        task = _make_task(
            tid="case-1",
            output_dir="artifact",
            rubric={"mandatory": ["criterion"]},
        )
        captured: dict[str, Any] = {}

        def fake_judge(*args: Any, **kwargs: Any) -> dict[str, Any]:
            captured.update(kwargs)
            return {"mandatory": [True], "good_to_have": [], "ideal": []}

        with patch("eval.core._resolve_output_dir", return_value=out_dir):
            with patch("eval.core.judge_rubric", side_effect=fake_judge):
                with patch("eval.core.score_rubric", return_value=0.4):
                    result = _judge_task_with_rubric(task, "ans", criterion_workers=1)

        assert result is not None
        assert captured["output_dir"] == str(out_dir)
        assert "artifact_context" in captured
        assert "summary.txt" in captured["artifact_context"]
        assert captured["repl_seed"]["artifact_root"] == str(out_dir)
        assert captured["repl_seed"]["artifacts_text"]["summary.txt"] == "hello judge"


# ---------------------------------------------------------------------------
# Resume: index preservation
# ---------------------------------------------------------------------------

class TestResumeIndexPreservation:
    """The bug was: after filtering completed tasks, remaining indices
    must match the original dataset positions so _to_case gets the right idx."""

    def test_remaining_preserves_original_indices(self, tmp_path: Path) -> None:
        from service.engine import _extract_completed, _task_id, _to_case, AsyncRunConfig

        tasks = [
            {"id": "a", "task": "first"},
            {"id": "b", "task": "second"},
            {"id": "c", "task": "third"},
        ]
        # Simulate "b" already completed
        out = tmp_path / "out.jsonl"
        out.write_text(json.dumps({"id": "b", "status": "ok"}) + "\n")
        completed_ids, _, _ = _extract_completed(out)

        remaining = [
            (idx, task)
            for idx, task in enumerate(tasks)
            if _task_id(task, idx) not in completed_ids
        ]
        remaining_tasks = [task for _, task in remaining]
        remaining_cases = [_to_case(task, idx, AsyncRunConfig()) for idx, task in remaining]

        assert len(remaining_cases) == 2
        assert remaining_cases[0].case_id == "a"
        assert remaining_cases[1].case_id == "c"
        # Verify task/case alignment
        assert remaining_tasks[0]["id"] == "a"
        assert remaining_tasks[1]["id"] == "c"


# ---------------------------------------------------------------------------
# Storage snapshot sync
# ---------------------------------------------------------------------------

class TestStorageSnapshotSync:
    def test_schedule_sync_creates_snapshot(self, tmp_path: Path) -> None:
        from service.engine import _schedule_sync_snapshot

        out = tmp_path / "out.jsonl"
        meta = tmp_path / "out.meta.json"
        out.write_text("line1\n")
        meta.write_text('{"status": "running"}')

        storage = MagicMock()
        # Mock fire_and_forget to prevent background thread from cleaning up
        # snapshots before we can assert on them.
        with patch("service.engine.fire_and_forget") as mock_ff:
            _schedule_sync_snapshot(storage, out, meta)

        # Snapshots created in .sync_tmp/
        sync_dir = tmp_path / ".sync_tmp"
        assert sync_dir.exists()
        snaps = list(sync_dir.iterdir())
        assert len(snaps) == 2
        # fire_and_forget was called with the upload function
        mock_ff.assert_called_once()

    def test_upload_snapshot_and_cleanup(self, tmp_path: Path) -> None:
        from service.engine import _upload_snapshot_and_cleanup

        snap_out = tmp_path / "out.snap"
        snap_meta = tmp_path / "meta.snap"
        snap_out.write_text("data")
        snap_meta.write_text("meta")

        storage = MagicMock()
        _upload_snapshot_and_cleanup(storage, "out.jsonl", snap_out, "out.meta.json", snap_meta)

        storage.put.assert_any_call("out.jsonl", snap_out)
        storage.put.assert_any_call("out.meta.json", snap_meta)
        # Snapshots cleaned up
        assert not snap_out.exists()
        assert not snap_meta.exists()

    def test_upload_cleanup_on_error(self, tmp_path: Path) -> None:
        from service.engine import _upload_snapshot_and_cleanup

        snap_out = tmp_path / "out.snap"
        snap_meta = tmp_path / "meta.snap"
        snap_out.write_text("data")
        snap_meta.write_text("meta")

        storage = MagicMock()
        storage.put.side_effect = RuntimeError("s3 down")
        # Should not raise — error is logged
        _upload_snapshot_and_cleanup(storage, "out.jsonl", snap_out, "out.meta.json", snap_meta)
        # Snapshots still cleaned up
        assert not snap_out.exists()
        assert not snap_meta.exists()


# ---------------------------------------------------------------------------
# PythonREPL (eval/tools.py)
# ---------------------------------------------------------------------------

class TestPythonREPL:
    def test_basic_execution(self) -> None:
        from eval.tools import PythonREPL

        r = PythonREPL()
        out = r.run("print(1 + 2)")
        assert "3" in out

    def test_state_persists(self) -> None:
        from eval.tools import PythonREPL

        r = PythonREPL()
        r.run("x = 42")
        out = r.run("print(x)")
        assert "42" in out

    def test_error_captured(self) -> None:
        from eval.tools import PythonREPL

        r = PythonREPL()
        out = r.run("raise ValueError('boom')")
        assert "ValueError" in out
        assert "boom" in out

    def test_reset_clears_state(self) -> None:
        from eval.tools import PythonREPL

        r = PythonREPL()
        r.run("x = 42")
        r.reset()
        out = r.run("print(x)")
        assert "NameError" in out

    def test_truncation(self) -> None:
        from eval.tools import PythonREPL, MAX_OUTPUT

        r = PythonREPL()
        out = r.run(f"print('x' * {MAX_OUTPUT + 1000})")
        assert len(out) <= MAX_OUTPUT + 100  # some slack for truncation message
        assert "TRUNCATED" in out

    def test_timeout_is_enforced(self) -> None:
        from eval.tools import PythonREPL

        r = PythonREPL()
        started = time.perf_counter()
        out = r.run("while True:\n    pass", timeout=1)
        elapsed = time.perf_counter() - started

        assert "[TIMEOUT after 1s]" in out
        assert elapsed < 3

    def test_repl_cwd_is_eval_dir(self) -> None:
        """REPL cwd must be eval/ so relative paths like 'artifacts/foo.xlsx' resolve to eval/artifacts/."""
        from eval.tools import PythonREPL
        from pathlib import Path

        r = PythonREPL()
        out = r.run("import os; print(os.getcwd())")
        repl_cwd = out.strip()
        eval_dir = str(Path(__file__).resolve().parent.parent / "eval")
        assert repl_cwd == eval_dir, f"REPL cwd {repl_cwd} != expected {eval_dir}"

    def test_resolve_output_dir_matches_repl_cwd(self) -> None:
        """_resolve_output_dir('artifacts') must point to eval/artifacts/ (same base as REPL cwd)."""
        from eval.core import _resolve_output_dir, EVAL_DIR

        resolved = _resolve_output_dir("artifacts")
        assert resolved == EVAL_DIR / "artifacts"

    def test_repl_can_read_file_in_output_dir(self, tmp_path: Path) -> None:
        """REPL should read files written to the output_dir resolved by _resolve_output_dir."""
        from eval.tools import PythonREPL
        from eval.core import _resolve_output_dir

        out_dir = _resolve_output_dir("artifacts")
        marker = tmp_path / "marker.txt"
        # Write a temp file inside eval/artifacts to verify the REPL finds it
        test_file = out_dir / "_test_repl_read.tmp"
        test_file.write_text("hello_from_test")
        try:
            r = PythonREPL()
            out = r.run("print(open('artifacts/_test_repl_read.tmp').read())")
            assert "hello_from_test" in out
        finally:
            test_file.unlink(missing_ok=True)

    def test_seed_globals_available(self) -> None:
        """PythonREPL(seed_globals=...) should make variables immediately available."""
        from eval.tools import PythonREPL

        r = PythonREPL(seed_globals={"my_var": 123})
        out = r.run("print(my_var)")
        assert "123" in out

    def test_output_dir_redirects_escaped_writes(self, tmp_path: Path) -> None:
        from eval.tools import PythonREPL

        out_dir = tmp_path / "artifact"
        r = PythonREPL(output_dir=str(out_dir))
        out = r.run("open('../escape.txt', 'w').write('hello'); print('done')")

        assert "done" in out
        assert not (tmp_path / "escape.txt").exists()
        assert (out_dir / "escape.txt").read_text() == "hello"

    def test_output_dir_nested_files_are_collected_for_judging(self, tmp_path: Path) -> None:
        from eval.core import _collect_output_files
        from eval.tools import PythonREPL

        out_dir = tmp_path / "artifact"
        r = PythonREPL(output_dir=str(out_dir))
        out = r.run("open('reports/final.txt', 'w').write('nested hello'); print('done')")

        collected = _collect_output_files(out_dir)

        assert "done" in out
        assert "reports/final.txt" in collected["inline"]
        assert "nested hello" in collected["inline"]


# ---------------------------------------------------------------------------
# Tool concurrency tracking (eval/tools.py)
# ---------------------------------------------------------------------------

class TestToolConcurrency:
    def test_set_and_get_metrics(self) -> None:
        from eval.tools import reset_tool_metrics, get_tool_metrics, set_tool_concurrency

        set_tool_concurrency(4)
        reset_tool_metrics()
        m = get_tool_metrics()
        assert m["in_flight_tools"] == 0
        assert m["peak_in_flight_tools"] == 0

    def test_set_concurrency_none_disables(self) -> None:
        from eval.tools import set_tool_concurrency
        import eval.tools as t

        set_tool_concurrency(None)
        assert t._tool_sem is None

    def test_set_concurrency_invalid(self) -> None:
        from eval.tools import set_tool_concurrency

        with pytest.raises(ValueError, match="must be >= 1"):
            set_tool_concurrency(0)


# ---------------------------------------------------------------------------
# on_case_complete callback in Runner
# ---------------------------------------------------------------------------

class TestOnCaseComplete:
    def test_callback_receives_all_results(self) -> None:
        from clients import FakeLLMClient

        client = FakeLLMClient(base_latency_s=0.001, jitter_s=0.001, tool_ratio=0.0)
        runner = __import__("runner").Runner(llm_client=client, eval_sem=4, cpu_sem=2)
        cases = [EvalCase(case_id=f"c-{i}", prompt="hi", max_steps=1) for i in range(10)]

        collected: list[CaseResult] = []

        async def _cb(result: CaseResult) -> None:
            collected.append(result)

        asyncio.run(runner.run_cases(cases, on_case_complete=_cb))
        assert len(collected) == 10
        assert {r.case_id for r in collected} == {f"c-{i}" for i in range(10)}

    def test_sync_callback_also_works(self) -> None:
        from clients import FakeLLMClient

        client = FakeLLMClient(base_latency_s=0.001, jitter_s=0.001, tool_ratio=0.0)
        runner = __import__("runner").Runner(llm_client=client, eval_sem=4, cpu_sem=2)
        cases = [EvalCase(case_id="s1", prompt="hi", max_steps=1)]

        collected: list[str] = []

        def _cb(result: CaseResult) -> None:
            collected.append(result.case_id)

        asyncio.run(runner.run_cases(cases, on_case_complete=_cb))
        assert collected == ["s1"]


class TestProviderClientConstruction:
    def test_load_provider_module_clears_stale_openai_env(self, monkeypatch: Any) -> None:
        import service.engine as engine_mod

        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        monkeypatch.setenv("OPENAI_MODEL", "nvidia/nemotron-3-super-120b-a12b")
        monkeypatch.setenv("OPENAI_REASONING_EFFORT", "xhigh")

        module = engine_mod._load_provider_module(
            engine_mod.BASE_DIR / "eval" / "llms" / "openai.py",
            {},
        )

        assert module.LLM_ID == "gpt-5.4"
        assert "OPENAI_MODEL" not in os.environ
        assert "OPENAI_REASONING_EFFORT" not in os.environ

    def test_provider_case_context_seeds_docx_reference_data(self, tmp_path: Path, monkeypatch: Any) -> None:
        import service.engine as engine_mod
        import eval.core as core_mod
        from docx import Document

        refs_dir = tmp_path / "reference_files"
        refs_dir.mkdir()
        docx_path = refs_dir / "quotes.docx"

        doc = Document()
        doc.add_paragraph("Example quotation")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Vendor"
        table.rows[0].cells[1].text = "Price"
        table.rows[1].cells[0].text = "Vendocrat"
        table.rows[1].cells[1].text = "123"
        doc.save(str(docx_path))

        monkeypatch.setattr(engine_mod, "BASE_DIR", tmp_path)
        monkeypatch.setattr(core_mod, "BASE_DIR", tmp_path)
        monkeypatch.setattr(
            engine_mod,
            "_prepare_task_output_dir",
            lambda task: str(tmp_path / "artifact" / str(task.get("id", ""))),
        )

        ctx = engine_mod._provider_case_context(
            {"id": "case-docx", "task": "analyze", "reference_files": ["quotes.docx"]}
        )

        repl = ctx.config["_repl"]
        seed = repl._seed_globals  # type: ignore[attr-defined]
        assert seed is not None
        assert "_reference_files" in seed
        assert "quotes.docx" in seed["_reference_files"]
        assert seed["_reference_files"]["quotes.docx"]["tables"][0][1] == ["Vendocrat", "123"]
        assert "quotes_tables" in seed
        assert "Do not try to reopen the original reference files from disk." in ctx.config["_prompt_repl_note"]
        assert "`quotes_tables`" in ctx.config["_prompt_repl_note"]

    def test_build_client_defers_provider_context_creation(self, monkeypatch: Any) -> None:
        import service.engine as engine_mod

        calls: list[str] = []

        def fake_provider_case_context(
            task: dict[str, Any], *, conv_store: Any = None
        ) -> Any:
            calls.append(str(task["id"]))
            return engine_mod.ProviderCaseContext(
                task_text=str(task["task"]),
                references={},
                config={},
            )

        class DemoProvider:
            LLM_ID = "demo-provider"

            @staticmethod
            def generate(
                task: str, references: dict[str, Any], config: dict[str, Any]
            ) -> dict[str, Any]:
                return {"text": task}

        monkeypatch.setattr(engine_mod, "_provider_case_context", fake_provider_case_context)
        monkeypatch.setattr(
            engine_mod,
            "_resolve_provider_module",
            lambda provider, provider_config_path: (DemoProvider, provider),
        )

        llm_id, client = engine_mod._build_client(
            engine_mod.AsyncRunConfig(client="provider", provider="demo"),
            task_by_id={
                "case-a": {"id": "case-a", "task": "Task A"},
                "case-b": {"id": "case-b", "task": "Task B"},
            },
        )

        assert llm_id == "demo-provider"
        assert calls == []

        async def _run_once() -> Any:
            return await client.complete(
                [{"role": "user", "content": "Task A", "case_id": "case-a", "step": 0}]
            )

        result = asyncio.run(_run_once())
        assert result.text == "Task A"
        assert calls == ["case-a"]


# ---------------------------------------------------------------------------
# max_steps exceeded
# ---------------------------------------------------------------------------

class TestMaxStepsExceeded:
    def test_error_when_tool_always_needed(self) -> None:
        from clients import LLMClient, Response

        class AlwaysToolClient(LLMClient):
            async def complete(self, messages: list[dict[str, Any]]) -> Response:
                return Response(text="tool-again", tool_needed=True, tool_input=0.001)

        runner = __import__("runner").Runner(
            llm_client=AlwaysToolClient(), eval_sem=2, cpu_sem=2
        )
        cases = [EvalCase(case_id="loop", prompt="go", tool_mode="sleep", max_steps=2)]
        results = asyncio.run(runner.run_cases(cases))

        assert len(results) == 1
        assert results[0].status == "error"
        assert "max_steps_exceeded" in results[0].error


# ---------------------------------------------------------------------------
# Empty case list
# ---------------------------------------------------------------------------

class TestEmptyCases:
    def test_runner_handles_empty(self) -> None:
        from clients import FakeLLMClient

        runner = __import__("runner").Runner(
            llm_client=FakeLLMClient(), eval_sem=4, cpu_sem=2
        )
        results = asyncio.run(runner.run_cases([]))
        assert results == []


# ---------------------------------------------------------------------------
# FakeLLMClient determinism
# ---------------------------------------------------------------------------

class TestFakeLLMClientDeterminism:
    def test_same_input_same_output(self) -> None:
        from clients import FakeLLMClient

        c1 = FakeLLMClient(base_latency_s=0.0, jitter_s=0.0, tool_ratio=0.5)
        c2 = FakeLLMClient(base_latency_s=0.0, jitter_s=0.0, tool_ratio=0.5)
        msgs = [{"role": "user", "content": "hi", "case_id": "det", "step": 0}]

        r1 = asyncio.run(c1.complete(msgs))
        r2 = asyncio.run(c2.complete(msgs))
        assert r1.text == r2.text
        assert r1.tool_needed == r2.tool_needed

    def test_force_tool(self) -> None:
        from clients import FakeLLMClient

        c = FakeLLMClient(force_tool=True)
        msgs = [{"role": "user", "content": "hi", "case_id": "ft", "step": 0}]
        r = asyncio.run(c.complete(msgs))
        assert r.tool_needed is True

    def test_no_tool_on_step_1(self) -> None:
        from clients import FakeLLMClient

        c = FakeLLMClient(force_tool=True)
        msgs = [{"role": "user", "content": "hi", "case_id": "ft", "step": 1}]
        r = asyncio.run(c.complete(msgs))
        assert r.tool_needed is False


# ---------------------------------------------------------------------------
# Full end-to-end: _run_async_eval with FakeLLMClient (no tokens)
# ---------------------------------------------------------------------------

class TestRunAsyncEvalE2E:
    def test_fake_client_e2e(self, tmp_path: Path) -> None:
        from service.engine import AsyncRunConfig, run_async_eval

        ds = tmp_path / "ds.jsonl"
        tasks = [{"id": f"e2e-{i}", "task": f"task {i}"} for i in range(8)]
        ds.write_text("\n".join(json.dumps(t) for t in tasks) + "\n")

        config = AsyncRunConfig(
            client="fake",
            eval_sem=4,
            cpu_sem=2,
            case_max_steps=2,
            fake_base_latency_s=0.001,
            fake_jitter_s=0.001,
        )
        out = run_async_eval(config, dataset_path=ds, output_path=tmp_path / "out.jsonl")

        assert out.exists()
        rows = [json.loads(line) for line in out.read_text().splitlines()]
        assert len(rows) == 8
        assert all(r["runner"] == "async" for r in rows)
        assert all(r["llm_id"] == "async-fake" for r in rows)

        meta = json.loads(out.with_suffix(".meta.json").read_text())
        assert meta["status"] == "completed"
        assert meta["completed"] == 8
        assert meta["ok"] == 8

    def test_fake_client_with_task_ids_filter(self, tmp_path: Path) -> None:
        from service.engine import AsyncRunConfig, run_async_eval

        ds = tmp_path / "ds.jsonl"
        tasks = [{"id": f"f-{i}", "task": f"task {i}"} for i in range(10)]
        ds.write_text("\n".join(json.dumps(t) for t in tasks) + "\n")

        config = AsyncRunConfig(
            client="fake",
            fake_base_latency_s=0.001,
            fake_jitter_s=0.001,
        )
        out = run_async_eval(
            config,
            dataset_path=ds,
            output_path=tmp_path / "out.jsonl",
            task_ids={"f-2", "f-5", "f-9"},
        )
        rows = [json.loads(line) for line in out.read_text().splitlines()]
        assert len(rows) == 3
        assert {r["id"] for r in rows} == {"f-2", "f-5", "f-9"}

    def test_resume_skips_completed(self, tmp_path: Path) -> None:
        from service.engine import AsyncRunConfig, run_async_eval

        ds = tmp_path / "ds.jsonl"
        tasks = [{"id": f"r-{i}", "task": f"task {i}"} for i in range(5)]
        ds.write_text("\n".join(json.dumps(t) for t in tasks) + "\n")

        out = tmp_path / "out.jsonl"
        # Pre-seed 2 completed results
        with out.open("w") as f:
            f.write(json.dumps({"id": "r-0", "status": "ok", "case_id": "r-0"}) + "\n")
            f.write(json.dumps({"id": "r-1", "status": "ok", "case_id": "r-1"}) + "\n")

        config = AsyncRunConfig(
            client="fake",
            fake_base_latency_s=0.001,
            fake_jitter_s=0.001,
        )
        run_async_eval(config, dataset_path=ds, output_path=out)

        rows = [json.loads(line) for line in out.read_text().splitlines()]
        assert len(rows) == 5  # 2 pre-existing + 3 new
        ids = [r["id"] for r in rows]
        # First 2 are the pre-seeded rows (no runner key)
        assert "runner" not in rows[0]
        # Last 3 are the new ones
        assert all(r["runner"] == "async" for r in rows[2:])

    def test_judge_enabled_with_mock(self, tmp_path: Path) -> None:
        from service.engine import AsyncRunConfig, run_async_eval

        ds = tmp_path / "ds.jsonl"
        tasks = [
            {
                "id": "j-0",
                "task": "solve it",
                "rubric": {"mandatory": ["is correct"], "good_to_have": [], "ideal": []},
            },
            {
                "id": "j-1",
                "task": "other task",
                # No rubric — judge should skip
            },
        ]
        ds.write_text("\n".join(json.dumps(t) for t in tasks) + "\n")

        mock_eval = {"mandatory": [True], "good_to_have": [], "ideal": [], "score": 0.4}
        config = AsyncRunConfig(
            client="fake",
            judge_enabled=True,
            judge_sem=2,
            judge_criterion_workers=1,
            fake_base_latency_s=0.001,
            fake_jitter_s=0.001,
        )

        with patch("service.engine._judge_task_with_rubric", return_value=mock_eval) as mock_fn:
            out = run_async_eval(config, dataset_path=ds, output_path=tmp_path / "out.jsonl")

        rows = [json.loads(line) for line in out.read_text().splitlines()]
        assert len(rows) == 2

        # The task with rubric should have eval
        judged = [r for r in rows if r["id"] == "j-0"]
        assert len(judged) == 1
        assert judged[0]["eval"]["score"] == 0.4

        meta = json.loads(out.with_suffix(".meta.json").read_text())
        assert meta["status"] == "completed"

    def test_replay_client_e2e(self, tmp_path: Path) -> None:
        from service.engine import AsyncRunConfig, run_async_eval

        ds = tmp_path / "ds.jsonl"
        tasks = [{"id": "rp-0", "task": "prompt"}, {"id": "rp-1", "task": "prompt2"}]
        ds.write_text("\n".join(json.dumps(t) for t in tasks) + "\n")

        fixtures = {
            "rp-0": {"0": {"text": "done-0", "tool_needed": False}},
            "rp-1": {"0": {"text": "done-1", "tool_needed": False}},
        }
        fix_path = tmp_path / "fixtures.json"
        fix_path.write_text(json.dumps(fixtures))

        config = AsyncRunConfig(
            client="replay",
            replay_fixtures=str(fix_path),
        )
        out = run_async_eval(config, dataset_path=ds, output_path=tmp_path / "out.jsonl")
        rows = [json.loads(line) for line in out.read_text().splitlines()]
        assert len(rows) == 2
        texts = {r["id"]: r["llm_answer"] for r in rows}
        assert texts["rp-0"] == "done-0"
        assert texts["rp-1"] == "done-1"


# ---------------------------------------------------------------------------
# API request validation (no server startup needed)
# ---------------------------------------------------------------------------

class TestApiValidation:
    def _make_req(self, **overrides: Any) -> Any:
        from service.api import EvalRequest
        defaults = {"engine": "async", "client": "fake"}
        defaults.update(overrides)
        return EvalRequest(**defaults)

    def test_eval_sem_zero_rejected(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException) as exc_info:
            _validate_request(self._make_req(eval_sem=0))
        assert exc_info.value.status_code == 400

    def test_cpu_sem_zero_rejected(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException):
            _validate_request(self._make_req(cpu_sem=0))

    def test_negative_max_retries_rejected(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException):
            _validate_request(self._make_req(max_retries=-1))

    def test_legacy_requires_llm(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException, match="llm is required"):
            _validate_request(self._make_req(engine="legacy"))

    def test_replay_requires_fixtures(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException, match="replay_fixtures"):
            _validate_request(self._make_req(client="replay"))

    def test_provider_requires_llm(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException, match="llm is required"):
            _validate_request(self._make_req(client="provider"))

    def test_judge_sem_zero_rejected(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException):
            _validate_request(self._make_req(judge_sem=0))

    def test_judge_criterion_workers_zero_rejected(self) -> None:
        from service.api import _validate_request
        from fastapi import HTTPException

        with pytest.raises(HTTPException):
            _validate_request(self._make_req(judge_criterion_workers=0))

    def test_valid_request_passes(self) -> None:
        from service.api import _validate_request

        _validate_request(self._make_req())  # Should not raise


# ---------------------------------------------------------------------------
# ConvStore
# ---------------------------------------------------------------------------

class TestConvStore:
    """Tests for the ConvStore in-memory conversation store."""

    def test_append_and_flush(self, tmp_path: Path) -> None:
        from service.engine import ConvStore

        cf = tmp_path / "test.conv.jsonl"
        store = ConvStore(cf)
        store.append("kw_001", {"turn": 0, "role": "system", "content": "hello"})
        store.append("kw_001", {"turn": 1, "role": "assistant", "content": "hi"})
        assert not cf.exists()
        store.flush()

        lines = [json.loads(l) for l in cf.read_text().splitlines()]
        assert len(lines) == 1
        assert lines[0]["task_id"] == "kw_001"
        assert len(lines[0]["conversation"]) == 2
        assert lines[0]["conversation"][0]["role"] == "system"
        assert lines[0]["conversation"][1]["role"] == "assistant"

    def test_multiple_tasks(self, tmp_path: Path) -> None:
        from service.engine import ConvStore

        cf = tmp_path / "test.conv.jsonl"
        store = ConvStore(cf)
        store.append("kw_001", {"turn": 0, "role": "user", "content": "task1"})
        store.append("kw_002", {"turn": 0, "role": "user", "content": "task2"})
        store.append("kw_001", {"turn": 1, "role": "assistant", "content": "done1"})
        store.flush()

        lines = [json.loads(l) for l in cf.read_text().splitlines()]
        assert len(lines) == 2
        by_id = {l["task_id"]: l for l in lines}
        assert len(by_id["kw_001"]["conversation"]) == 2
        assert len(by_id["kw_002"]["conversation"]) == 1

    def test_append_judge(self, tmp_path: Path) -> None:
        from service.engine import ConvStore

        cf = tmp_path / "test.conv.jsonl"
        store = ConvStore(cf)
        store.append("kw_001", {"turn": 0, "role": "user", "content": "task"})
        store.append_judge("kw_001", {"tier": "mandatory", "verdict": "PASS"})
        store.flush()

        lines = [json.loads(l) for l in cf.read_text().splitlines()]
        assert len(lines) == 1
        assert len(lines[0]["conversation"]) == 1
        assert len(lines[0]["judge"]) == 1
        assert lines[0]["judge"][0]["verdict"] == "PASS"

    def test_reload_uses_latest_snapshot_for_task(self, tmp_path: Path) -> None:
        from service.engine import ConvStore

        cf = tmp_path / "test.conv.jsonl"
        store = ConvStore(cf)
        store.append("kw_001", {"turn": 0, "role": "user", "content": "task"})
        store.flush()

        store.append("kw_001", {"turn": 1, "role": "assistant", "content": "done"})
        store.append_judge("kw_001", {"tier": "mandatory", "verdict": "PASS"})
        store.flush()

        lines = [json.loads(l) for l in cf.read_text().splitlines()]
        assert len(lines) == 2

        reloaded = ConvStore(cf)
        reloaded.append("kw_001", {"turn": 2, "role": "assistant", "content": "again"})
        reloaded.flush()
        latest = [json.loads(l) for l in cf.read_text().splitlines()][-1]
        assert latest["task_id"] == "kw_001"
        assert len(latest["conversation"]) == 3
        assert len(latest["judge"]) == 1

    def test_judge_task_passes_conv_store(self) -> None:
        from service.engine import _judge_task_with_rubric, ConvStore

        task = _make_task("t1", "solve", rubric={
            "mandatory": ["m1"],
            "good_to_have": [],
            "ideal": [],
        })
        mock_eval = {"mandatory": [True], "good_to_have": [], "ideal": []}

        captured = {}

        def mock_judge_rubric(*args, **kwargs):
            captured.update(kwargs)
            return mock_eval

        with patch("eval.core.judge_rubric", side_effect=mock_judge_rubric):
            with patch("eval.core.score_rubric", return_value=0.4):
                _judge_task_with_rubric(
                    task, "answer", criterion_workers=1,
                    conv_store=MagicMock(),
                )

        assert captured["conv_store"] is not None
        assert captured["task_id"] == "t1"
