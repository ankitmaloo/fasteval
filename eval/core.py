"""Core eval logic: judge rubric criteria, compute scores, run eval loop."""

from __future__ import annotations

import json
import os
import csv
import re
import shutil
import asyncio
import time
import threading
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import TYPE_CHECKING, Any

from google import genai
from google.genai import types
from dotenv import load_dotenv

from eval.tools import CODE_SCHEMA, PythonREPL, make_repl, repl
from eval.log import log

try:
    import weave
    _weave_op = weave.op
except ImportError:
    def _weave_op(fn=None, **kwargs):
        return fn if fn is not None else lambda f: f

if TYPE_CHECKING:
    from eval.storage import Storage

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent
EVAL_DIR = Path(__file__).resolve().parent
DATASET_PATH = BASE_DIR / "dataset.jsonl"
RESULTS_DIR = EVAL_DIR / "results"

JUDGE_PROMPT = """You are a strict evaluator. You will be given:
- A TASK that was given to an AI agent
- The agent's ANSWER
- An ARTIFACT CONTEXT block describing the task artifact directory, discovered files, and preloaded objects
- A single CRITERION to evaluate

You have a Python REPL tool with the task artifact directory as its working directory when one exists.
Use `print(artifact_root)`, `print(artifact_manifest)`, and `print(_preloaded_files)` to see what was discovered and preloaded.
Common pre-loaded variables:
- `artifact_root`: absolute artifact directory for this task
- `artifact_manifest`: list of discovered artifact metadata
- `artifacts_text`: parsed text artifacts keyed by relative path
- `artifacts_tables`: parsed table artifacts keyed by relative path
- `<filename>_data`: openpyxl Workbook opened with data_only=True (cached computed values)
- `<filename>_formulas`: openpyxl Workbook opened with data_only=False (raw formulas)
Example: `ws = model_i_headlamp_vendor_npv_analysis_data['Sheet1']` or `print(artifacts_text['report.md'])`.

Use the REPL for:
- Inspecting pre-loaded artifact objects
- Inspecting workbook cells and sheets
- Validating numerical calculations or formulas
- Counting words or characters precisely

Prefer the ARTIFACT CONTEXT block and preloaded variables over rediscovering file paths.

Respond with exactly one word: PASS or FAIL

No explanation, no hedging, just PASS or FAIL."""


@dataclass(frozen=True, slots=True)
class JudgeRuntime:
    provider: str
    kind: str
    model: str
    provider_config_path: str | None = None
    base_url: str | None = None
    api_key: str | None = None
    reasoning_effort: str | None = None
    extra_body: dict[str, Any] | None = None


_JUDGE_OPENAI_TOOL = {"type": "function", **CODE_SCHEMA}
_JUDGE_OPENAI_CHAT_TOOL = {"type": "function", "function": CODE_SCHEMA}
_JUDGE_ANTHROPIC_TOOL = {
    "name": CODE_SCHEMA["name"],
    "description": CODE_SCHEMA["description"],
    "input_schema": CODE_SCHEMA["parameters"],
}


def _load_eval_config(config_path: str | None = None) -> dict[str, Any]:
    try:
        import yaml  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyYAML is required to read eval/config.yaml") from exc
    path = Path(config_path) if config_path else (EVAL_DIR / "config.yaml")
    if not path.exists():
        return {}
    payload = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    if not isinstance(payload, dict):
        raise ValueError(f"Invalid config payload in {path}")
    return payload


def _resolve_judge_provider_name(
    judge_provider: str | None = None,
    provider_config_path: str | None = None,
) -> str:
    if judge_provider:
        return judge_provider
    payload = _load_eval_config(provider_config_path)
    judge_cfg = payload.get("judge", {})
    if isinstance(judge_cfg, dict) and judge_cfg.get("provider"):
        return str(judge_cfg["provider"])
    return "gemini"


@lru_cache(maxsize=32)
def _resolve_judge_runtime(
    judge_provider: str | None = None,
    provider_config_path: str | None = None,
) -> JudgeRuntime:
    from service.engine import _resolve_provider_module

    provider_name = _resolve_judge_provider_name(judge_provider, provider_config_path)
    module, _ = _resolve_provider_module(provider_name, provider_config_path)
    stem = Path(getattr(module, "__file__", "")).stem

    if stem == "gemini":
        return JudgeRuntime(
            provider=provider_name,
            kind="gemini",
            model=os.environ.get("GEMINI_MODEL", getattr(module, "LLM_ID", "gemini-3-flash-preview")),
            provider_config_path=provider_config_path,
        )
    if stem == "openai":
        return JudgeRuntime(
            provider=provider_name,
            kind="openai_responses",
            model=os.environ.get("OPENAI_MODEL", getattr(module, "LLM_ID", "gpt-5.4")),
            provider_config_path=provider_config_path,
            api_key=os.environ.get("OPENAI_API_KEY"),
            reasoning_effort=os.environ.get("OPENAI_REASONING_EFFORT", "medium"),
        )
    if stem == "oaichat":
        extra_body_raw = os.environ.get("OAICHAT_EXTRA_BODY")
        return JudgeRuntime(
            provider=provider_name,
            kind="openai_chat",
            model=os.environ["OAICHAT_MODEL"],
            provider_config_path=provider_config_path,
            base_url=os.environ["OAICHAT_BASE_URL"],
            api_key=os.environ["OAICHAT_API_KEY"],
            extra_body=json.loads(extra_body_raw) if extra_body_raw else None,
        )
    if stem == "claude":
        return JudgeRuntime(
            provider=provider_name,
            kind="anthropic",
            model=getattr(module, "LLM_ID", "claude-opus-4-6"),
            provider_config_path=provider_config_path,
            api_key=os.environ.get("ANTHROPIC_API_KEY"),
        )
    if stem == "ant_compat":
        return JudgeRuntime(
            provider=provider_name,
            kind="anthropic",
            model=os.environ["ANTCOMPAT_MODEL"],
            provider_config_path=provider_config_path,
            base_url=os.environ["ANTCOMPAT_BASE_URL"],
            api_key=os.environ["ANTCOMPAT_API_KEY"],
        )
    raise ValueError(f"Unsupported judge provider backend for '{provider_name}': {stem}")


def _judge_usage_dict(usage: Any) -> dict[str, int]:
    if usage is None:
        return {}
    return {
        "input_tokens": int(
            getattr(usage, "prompt_token_count", None)
            or getattr(usage, "input_tokens", 0)
            or 0
        ),
        "output_tokens": int(
            getattr(usage, "candidates_token_count", None)
            or getattr(usage, "output_tokens", 0)
            or 0
        ),
    }


def load_dataset(path: Path | None = None) -> list[dict]:
    p = path or DATASET_PATH
    return [json.loads(line) for line in open(p)]


# Text extensions → inline as-is
_INLINE_EXTS = {".md", ".txt", ".json", ".jsonl", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".log"}


def _parse_docx(path: Path) -> str:
    """Extract all text (paragraphs + tables) from a .docx file."""
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        parts.append(f"\n[Table {i}]\n" + "\n".join(rows))
    return "\n".join(parts)


def _parse_xlsx(path: Path) -> str:
    """Extract all cell values from an Excel workbook as text.
    Uses data_only=True for cached formula values, falls back to formulas."""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    wb_formulas = openpyxl.load_workbook(str(path), data_only=False)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        ws_f = wb_formulas[sheet_name]
        rows = []
        for row_vals, row_formulas in zip(
            ws.iter_rows(values_only=False), ws_f.iter_rows(values_only=False)
        ):
            cells = []
            for cv, cf in zip(row_vals, row_formulas):
                v = cv.value
                if v is None and cf.value is not None:
                    v = cf.value  # fall back to formula string
                cells.append("" if v is None else str(v))
            rows.append(" | ".join(cells))
        parts.append(f"### Sheet: {sheet_name}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _parse_docx_tables(path: Path) -> list[list[list[str]]]:
    from docx import Document

    doc = Document(str(path))
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables


def read_reference_files(ref_paths: list[str] | None) -> dict:
    """Process reference files. Returns {inline: {name: content}, paths: {name: abs_path}}.
    Text/csv/md → inlined. Docx → parsed and inlined. Excel/binary → abs path for REPL."""
    if not ref_paths:
        return {"inline": {}, "paths": {}}
    inline = {}
    paths = {}
    for rel_path in ref_paths:
        full = BASE_DIR / rel_path
        if not full.exists():
            full = BASE_DIR / "reference_files" / rel_path
        if not full.exists():
            inline[rel_path] = f"[FILE NOT FOUND: {rel_path}]"
            continue
        ext = full.suffix.lower()
        if ext == ".docx":
            try:
                inline[rel_path] = _parse_docx(full)
            except Exception as e:
                inline[rel_path] = f"[DOCX PARSE ERROR: {e}]"
        elif ext in _INLINE_EXTS:
            try:
                inline[rel_path] = full.read_text(encoding="utf-8")
            except Exception as e:
                inline[rel_path] = f"[READ ERROR: {e}]"
        else:
            # binary, unknown → path for REPL
            paths[rel_path] = str(full.resolve())
    return {"inline": inline, "paths": paths}


_judge_code = types.FunctionDeclaration(
    name="execute_code",
    description=(
        "Execute Python code in a stateful REPL to verify claims. "
        "Variables persist across calls. Available: pandas, numpy, openpyxl, requests. "
        "Use for: counting words, validating calculations, checking URLs (requests.head), running code snippets."
    ),
    parameters=types.Schema(
        type="OBJECT",
        properties={"code": types.Schema(type="STRING", description="Python code to execute")},
        required=["code"],
    ),
)
_judge_tools = [types.Tool(function_declarations=[_judge_code])]

JUDGE_MAX_TURNS = 25


def _preload_output_files(output_dir: str | None, task_id: str | None = None) -> dict:
    """Backwards-compatible wrapper for judge REPL seeding."""
    return build_judge_artifact_bundle(output_dir, task_id=task_id)["repl_seed"]


def _artifact_var_base(rel_path: str) -> str:
    base = rel_path.lower()
    if "." in base:
        base = str(Path(base).with_suffix(""))
    return re.sub(r"[^a-z0-9]+", "_", base).strip("_") or "artifact"


def _format_manifest_entry(entry: dict[str, Any]) -> str:
    ext = entry.get("extension") or "[no ext]"
    preload = entry.get("preloaded")
    preload_text = ""
    if isinstance(preload, list) and preload:
        preload_text = "; preloaded: " + ", ".join(f"`{name}`" for name in preload)
    return (
        f"- {entry['relative_path']} ({ext}, {entry['size_bytes']} bytes"
        f"{preload_text})"
    )


def build_judge_artifact_bundle(output_dir: str | None, task_id: str | None = None) -> dict[str, Any]:
    """Build the deterministic artifact bundle consumed by judge prompt + REPL."""
    if not output_dir:
        return {
            "artifact_root": None,
            "manifest": [],
            "prompt_context": "No artifact directory was configured for this task.",
            "repl_seed": {},
        }

    import openpyxl

    artifact_dir = _resolve_output_dir(output_dir, task_id=task_id)
    artifact_root = str(artifact_dir.resolve())
    manifest: list[dict[str, Any]] = []
    repl_seed: dict[str, Any] = {
        "artifact_root": artifact_root,
        "artifact_manifest": manifest,
    }
    artifacts_text: dict[str, str] = {}
    artifacts_tables: dict[str, Any] = {}
    preloaded_files: dict[str, list[str]] = {}
    inline_sections: list[str] = []

    for path in _iter_output_files(artifact_dir):
        rel_path = path.relative_to(artifact_dir).as_posix()
        ext = path.suffix.lower()
        entry: dict[str, Any] = {
            "relative_path": rel_path,
            "absolute_path": str(path.resolve()),
            "extension": ext,
            "size_bytes": path.stat().st_size,
            "preloaded": [],
        }
        var_base = _artifact_var_base(rel_path)

        if ext in _INLINE_EXTS:
            try:
                content = path.read_text(encoding="utf-8")
            except Exception as exc:
                content = f"[READ ERROR: {exc}]"
            artifacts_text[rel_path] = content
            repl_seed[f"{var_base}_text"] = content
            entry["preloaded"].append(f"{var_base}_text")
            inline_sections.append(f"--- {rel_path} ---\n{content}")
        elif ext == ".docx":
            try:
                content = _parse_docx(path)
                tables = _parse_docx_tables(path)
            except Exception as exc:
                content = f"[DOCX PARSE ERROR: {exc}]"
                tables = []
            artifacts_text[rel_path] = content
            artifacts_tables[rel_path] = tables
            repl_seed[f"{var_base}_text"] = content
            repl_seed[f"{var_base}_tables"] = tables
            entry["preloaded"].extend([f"{var_base}_text", f"{var_base}_tables"])
            inline_sections.append(f"--- {rel_path} ---\n{content}")
        elif ext in (".xlsx", ".xls"):
            try:
                wb_data = openpyxl.load_workbook(str(path), data_only=True)
                wb_formulas = openpyxl.load_workbook(str(path), data_only=False)
                repl_seed[f"{var_base}_data"] = wb_data
                repl_seed[f"{var_base}_formulas"] = wb_formulas
                entry["preloaded"].extend([f"{var_base}_data", f"{var_base}_formulas"])
            except Exception as exc:
                log.warning(f"  [PRELOAD] failed to load {rel_path}: {exc}")

        if entry["preloaded"]:
            preloaded_files[rel_path] = list(entry["preloaded"])
        manifest.append(entry)

    if artifacts_text:
        repl_seed["artifacts_text"] = artifacts_text
    if artifacts_tables:
        repl_seed["artifacts_tables"] = artifacts_tables
    if preloaded_files:
        repl_seed["_preloaded_files"] = preloaded_files

    prompt_parts = [f"Artifact root: `{artifact_root}`"]
    if manifest:
        prompt_parts.append("Artifact manifest:")
        prompt_parts.extend(_format_manifest_entry(entry) for entry in manifest)
    else:
        prompt_parts.append("Artifact manifest: no files were produced. The model generated no output files. Do not attempt to access preloaded variables — none exist.")
    if preloaded_files:
        prompt_parts.append("Preloaded artifact objects:")
        for rel_path, names in preloaded_files.items():
            prompt_parts.append(
                f"- {rel_path}: " + ", ".join(f"`{name}`" for name in names)
            )
    if inline_sections:
        prompt_parts.append("--- ARTIFACT FILE CONTENTS ---")
        prompt_parts.append("\n\n".join(inline_sections))

    return {
        "artifact_root": artifact_root,
        "manifest": manifest,
        "prompt_context": "\n".join(prompt_parts),
        "repl_seed": repl_seed,
    }


@_weave_op
def _gemini_call_with_retry(client, model, contents, config, tag: str, max_retries: int = 6):
    """Call Gemini with retry on 429/5xx. Preserves conversation state."""
    import re
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(
                model=model, contents=contents, config=config,
            )
        except Exception as e:
            err_str = str(e)
            is_rate_limit = "429" in err_str or "RESOURCE_EXHAUSTED" in err_str
            is_server_err = any(c in err_str for c in ("500", "503", "UNAVAILABLE"))
            if not (is_rate_limit or is_server_err) or attempt == max_retries - 1:
                raise
            # Parse Gemini's suggested retry delay
            delay = 2 ** (attempt + 1)  # 2, 4, 8, 16, 32s
            m = re.search(r"retry in ([\d.]+)s", err_str, re.IGNORECASE)
            if m:
                delay = max(delay, float(m.group(1)) + 1)
            log.warning(f"  {tag} 429/5xx retry {attempt + 1}/{max_retries} (wait {delay:.0f}s)")
            time.sleep(delay)
    raise RuntimeError(f"{tag} exhausted {max_retries} retries")  # unreachable


def _judge_openai_responses(runtime: JudgeRuntime, prompt: str, judge_repl: Any, tag: str, _jconv) -> tuple[bool, str | None]:
    from openai import OpenAI

    client = OpenAI(api_key=runtime.api_key, timeout=6000, max_retries=0)
    input_items: list[dict[str, Any]] = [{"role": "user", "content": prompt}]

    for turn in range(JUDGE_MAX_TURNS):
        resp = client.responses.create(
            model=runtime.model,
            input=input_items,
            instructions=JUDGE_PROMPT,
            reasoning={"effort": runtime.reasoning_effort or "low"},
            tools=[_JUDGE_OPENAI_TOOL],
        )
        usage = _judge_usage_dict(resp.usage)
        fn_calls = [item for item in resp.output if item.type == "function_call"]
        if not fn_calls:
            verdict = (resp.output_text or "").strip().upper()
            passed = verdict.startswith("PASS")
            _jconv({
                "turn": turn + 1,
                "role": "judge_verdict",
                "verdict": "PASS" if passed else "FAIL",
                "response": resp.output_text or "",
                "usage": usage,
            })
            return passed, None

        input_items.extend(resp.output)
        for fc in fn_calls:
            code = json.loads(fc.arguments).get("code", "")
            result = judge_repl.run(code)
            _jconv({
                "turn": turn + 1,
                "role": "judge_tool",
                "tool": fc.name,
                "code": code,
                "result": result,
            })
            input_items.append({
                "type": "function_call_output",
                "call_id": fc.call_id,
                "output": result,
            })
    return False, None


def _judge_openai_chat(runtime: JudgeRuntime, prompt: str, judge_repl: Any, tag: str, _jconv) -> tuple[bool, str | None]:
    from openai import OpenAI

    client = OpenAI(base_url=runtime.base_url, api_key=runtime.api_key)
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": JUDGE_PROMPT},
        {"role": "user", "content": prompt},
    ]

    for turn in range(JUDGE_MAX_TURNS):
        kwargs: dict[str, Any] = {
            "model": runtime.model,
            "messages": messages,
            "tools": [_JUDGE_OPENAI_CHAT_TOOL],
            "tool_choice": "auto",
        }
        if runtime.extra_body:
            kwargs["extra_body"] = runtime.extra_body
        resp = client.chat.completions.create(**kwargs)
        raw = resp.model_dump()
        choice = raw["choices"][0]
        msg = choice["message"]
        tool_calls = msg.get("tool_calls") or []
        if not tool_calls:
            verdict = (msg.get("content") or "").strip().upper()
            passed = verdict.startswith("PASS")
            _jconv({
                "turn": turn + 1,
                "role": "judge_verdict",
                "verdict": "PASS" if passed else "FAIL",
                "response": msg.get("content") or "",
                "usage": raw.get("usage") or {},
            })
            return passed, None

        messages.append(msg)
        for tc in tool_calls:
            args = json.loads(tc["function"]["arguments"])
            code = args.get("code", "")
            result = judge_repl.run(code)
            _jconv({
                "turn": turn + 1,
                "role": "judge_tool",
                "tool": tc["function"]["name"],
                "code": code,
                "result": result,
            })
            messages.append({"role": "tool", "tool_call_id": tc["id"], "content": result})
    return False, None


def _judge_anthropic(runtime: JudgeRuntime, prompt: str, judge_repl: Any, tag: str, _jconv) -> tuple[bool, str | None]:
    import anthropic

    if runtime.base_url:
        client = anthropic.Anthropic(base_url=runtime.base_url, api_key=runtime.api_key)
    else:
        client = anthropic.Anthropic(api_key=runtime.api_key)
    messages: list[dict[str, Any]] = [{"role": "user", "content": prompt}]

    for turn in range(JUDGE_MAX_TURNS):
        resp = client.messages.create(
            model=runtime.model,
            max_tokens=4096,
            system=JUDGE_PROMPT,
            messages=messages,
            tools=[_JUDGE_ANTHROPIC_TOOL],
        )
        usage = _judge_usage_dict(resp.usage)
        tool_uses = [block for block in resp.content if block.type == "tool_use"]
        if not tool_uses:
            text = "".join(block.text for block in resp.content if getattr(block, "type", "") == "text")
            verdict = text.strip().upper()
            passed = verdict.startswith("PASS")
            _jconv({
                "turn": turn + 1,
                "role": "judge_verdict",
                "verdict": "PASS" if passed else "FAIL",
                "response": text,
                "usage": usage,
            })
            return passed, None

        messages.append({"role": "assistant", "content": resp.content})
        tool_results = []
        for tu in tool_uses:
            code = tu.input.get("code", "")
            result = judge_repl.run(code)
            _jconv({
                "turn": turn + 1,
                "role": "judge_tool",
                "tool": tu.name,
                "code": code,
                "result": result,
            })
            tool_results.append({"type": "tool_result", "tool_use_id": tu.id, "content": result})
        messages.append({"role": "user", "content": tool_results})
    return False, None


def _judge_gemini(runtime: JudgeRuntime, prompt: str, judge_repl: Any, tag: str, _jconv) -> tuple[bool, str | None]:
    client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
    contents = [types.Content(role="user", parts=[types.Part(text=prompt)])]
    config = types.GenerateContentConfig(
        system_instruction=JUDGE_PROMPT,
        thinking_config=types.ThinkingConfig(thinking_level="LOW"),
        tools=_judge_tools,
    )

    for turn in range(JUDGE_MAX_TURNS):
        resp = _gemini_call_with_retry(client, runtime.model, contents, config, tag)
        candidate = resp.candidates[0]
        usage = _judge_usage_dict(resp.usage_metadata)
        fn_calls = [p for p in candidate.content.parts if p.function_call]

        if not fn_calls:
            verdict = (resp.text or "").strip().upper()
            passed = verdict.startswith("PASS")
            _jconv({
                "turn": turn + 1,
                "role": "judge_verdict",
                "verdict": "PASS" if passed else "FAIL",
                "response": resp.text or "",
                "usage": usage,
            })
            return passed, None

        contents.append(candidate.content)
        fn_parts = []
        for part in fn_calls:
            fc = part.function_call
            code = fc.args.get("code", "")
            result = judge_repl.run(code)
            _jconv({
                "turn": turn + 1,
                "role": "judge_tool",
                "tool": fc.name,
                "code": code,
                "result": result,
            })
            fn_parts.append(types.Part(function_response=types.FunctionResponse(
                name=fc.name, response={"output": result},
            )))
        contents.append(types.Content(role="user", parts=fn_parts))
    return False, None


@_weave_op
def judge_criterion(task: str, answer: str, criterion: str, tier: str = "", idx: int = 0,
                    repl_seed: dict | None = None,
                    output_dir: str | None = None,
                    artifact_context: str | None = None,
                    conv_store: Any = None, task_id: str = "",
                    judge_provider: str | None = None,
                    provider_config_path: str | None = None,
                    repl_mode: str = "local",
                    sandbox_template: str | None = None) -> bool:
    """Judge a single criterion with its own REPL instance. Returns True for PASS."""
    tag = f"[JUDGE {tier}#{idx}] [{task_id}]"
    log.info(f"  {tag} evaluating: {criterion}")
    runtime = _resolve_judge_runtime(judge_provider, provider_config_path)
    remote_output_dir = None
    judge_seed = dict(repl_seed or {})
    judge_artifact_context = artifact_context
    task_slug = re.sub(r"[^A-Za-z0-9_.-]+", "-", task_id).strip("-") or "judge"
    if repl_mode == "daytona":
        remote_output_dir = f".kwbench/judge/{task_slug}/{tier}-{idx}"
        if judge_seed.get("artifact_root"):
            judge_seed["artifact_root"] = remote_output_dir
        manifest = judge_seed.get("artifact_manifest")
        if isinstance(manifest, list):
            patched_manifest = []
            for entry in manifest:
                if isinstance(entry, dict):
                    patched = dict(entry)
                    rel_path = patched.get("relative_path")
                    if isinstance(rel_path, str):
                        patched["absolute_path"] = f"{remote_output_dir}/{rel_path}"
                    patched_manifest.append(patched)
                else:
                    patched_manifest.append(entry)
            judge_seed["artifact_manifest"] = patched_manifest
        if artifact_context and output_dir:
            judge_artifact_context = artifact_context.replace(str(output_dir), remote_output_dir)
    judge_repl = make_repl(
        repl_mode=repl_mode,
        task_id=f"{task_slug}-{tier}-{idx}",
        seed_globals=judge_seed or None,
        output_dir=output_dir,
        remote_output_dir=remote_output_dir,
        sandbox_template=sandbox_template,
        sync_local_output_on_bootstrap=bool(output_dir),
    )
    prompt = (
        f"TASK:\n{task}\n\n"
        f"ANSWER:\n{answer}\n\n"
        f"ARTIFACT CONTEXT:\n{judge_artifact_context or 'No artifact context available.'}\n\n"
        f"CRITERION:\n{criterion}"
    )

    def _jconv(entry: dict) -> None:
        if conv_store and task_id:
            conv_store.append_judge(task_id, entry)

    _jconv({
        "turn": 0,
        "role": "judge_start",
        "provider": runtime.provider,
        "model": runtime.model,
        "tier": tier,
        "idx": idx,
        "criterion": criterion,
        "artifact_context": judge_artifact_context,
    })

    try:
        if runtime.kind == "gemini":
            passed, error = _judge_gemini(runtime, prompt, judge_repl, tag, _jconv)
        elif runtime.kind == "openai_responses":
            passed, error = _judge_openai_responses(runtime, prompt, judge_repl, tag, _jconv)
        elif runtime.kind == "openai_chat":
            passed, error = _judge_openai_chat(runtime, prompt, judge_repl, tag, _jconv)
        elif runtime.kind == "anthropic":
            passed, error = _judge_anthropic(runtime, prompt, judge_repl, tag, _jconv)
        else:
            raise ValueError(f"Unsupported judge runtime kind: {runtime.kind}")
        log.info(f"  {tag} -> {'PASS' if passed else 'FAIL'}")
        return passed, error
    except Exception as e:
        log.error(f"  {tag} FATAL: {e}")
        _jconv({"turn": -1, "role": "judge_error", "tier": tier, "idx": idx,
                "criterion": criterion, "error": str(e)})
        return (False, str(e))
    finally:
        try:
            judge_repl.close()
        except Exception:
            pass


@_weave_op
def judge_rubric(
    task: str,
    answer: str,
    rubric: dict,
    criterion_workers: int | None = None,
    repl_seed: dict | None = None,
    output_dir: str | None = None,
    artifact_context: str | None = None,
    conv_store: Any = None,
    task_id: str = "",
    judge_provider: str | None = None,
    provider_config_path: str | None = None,
    repl_mode: str = "local",
    sandbox_template: str | None = None,
) -> dict:
    """Judge all rubric criteria in parallel. Returns eval dict."""
    mandatory = rubric.get("mandatory", [])
    good_to_have = rubric.get("good_to_have", [])
    ideal = rubric.get("ideal", [])

    all_criteria = (
        [("mandatory", i, c) for i, c in enumerate(mandatory)]
        + [("good_to_have", i, c) for i, c in enumerate(good_to_have)]
        + [("ideal", i, c) for i, c in enumerate(ideal)]
    )

    results = {"mandatory": [None] * len(mandatory), "good_to_have": [None] * len(good_to_have), "ideal": [None] * len(ideal)}
    errors = []
    if not all_criteria:
        return results

    max_workers = min(16, len(all_criteria))
    if criterion_workers is not None:
        max_workers = max(1, min(max_workers, criterion_workers))

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(
                judge_criterion,
                task,
                answer,
                criterion,
                tier,
                idx,
                repl_seed,
                output_dir,
                artifact_context,
                conv_store=conv_store,
                task_id=task_id,
                judge_provider=judge_provider,
                provider_config_path=provider_config_path,
                repl_mode=repl_mode,
                sandbox_template=sandbox_template,
            ): (tier, idx)
            for tier, idx, criterion in all_criteria
        }
        for future in as_completed(futures):
            tier, idx = futures[future]
            passed, error = future.result()
            results[tier][idx] = passed
            if error:
                errors.append({"tier": tier, "idx": idx, "error": error})

    if errors:
        results["judge_errors"] = errors
        log.warning(f"  [JUDGE] {len(errors)} criterion(s) failed with errors")

    return results


def score_rubric(mandatory: list[bool], good_to_have: list[bool], ideal: list[bool]) -> float:
    if not all(mandatory):
        return 0.0
    score = 0.40
    if good_to_have:
        score += 0.35 * (sum(good_to_have) / len(good_to_have))
    if ideal:
        score += 0.25 * (sum(ideal) / len(ideal))
    return round(score, 4)


def build_prompt(
    task_text: str,
    references: dict,
    output_file: str | None = None,
    extra_instructions: str | None = None,
) -> str:
    """Build the full prompt string from task + references.
    references = {inline: {name: content}, paths: {name: abs_path}}"""
    parts = [task_text]
    inline = references.get("inline", {})
    file_paths = references.get("paths", {})
    if extra_instructions:
        parts.append(f"\n\n{extra_instructions}")
    if inline:
        parts.append("\n\n--- REFERENCE DOCUMENTS ---")
        for name, content in inline.items():
            parts.append(f"\n### {name}\n{content}")
    if file_paths:
        parts.append("\n\n--- DATA FILES (read these from disk using pandas or appropriate tools) ---")
        for name, abspath in file_paths.items():
            parts.append(f"- {name}: {abspath}")
    if output_file:
        parts.append(
            f"\n\nSave all output files only inside this path: `{output_file}`"
        )
    return "\n".join(parts)


def _load_completed_ids(path: Path) -> set[str]:
    """Read existing JSONL output, return set of completed task IDs."""
    if not path.exists():
        return set()
    ids = set()
    for line in open(path):
        try:
            ids.add(json.loads(line)["id"])
        except (json.JSONDecodeError, KeyError):
            continue
    return ids


def _write_meta(path: Path, **updates):
    """Atomic update of run metadata file."""
    meta = {}
    if path.exists():
        try:
            meta = json.loads(path.read_text())
        except (json.JSONDecodeError, OSError):
            pass
    meta.update(updates)
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(meta))
    tmp.rename(path)


def _resolve_output_dir(output_dir: str, task_id: str | None = None) -> Path:
    """Resolve task output_dir (from dataset) to absolute path relative to EVAL_DIR (matches REPL cwd)."""
    d = EVAL_DIR / output_dir
    if task_id:
        d = d / task_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _clear_output_dir(d: Path):
    """Remove all files in the output directory before a task run."""
    if not d.is_dir():
        return
    for child in d.iterdir():
        if child.is_dir():
            shutil.rmtree(child)
        elif child.is_file():
            child.unlink()


def _iter_output_files(d: Path) -> list[Path]:
    if not d.is_dir():
        return []
    return sorted(
        (path for path in d.rglob("*") if path.is_file()),
        key=lambda path: path.relative_to(d).as_posix(),
    )


def _collect_output_files(d: Path) -> dict:
    """Collect output files. Returns {inline: str, paths: list[str]}.
    Text files inlined, data files as absolute paths for REPL."""
    if not d.is_dir():
        return {"inline": "", "paths": []}
    inline_parts = []
    data_paths = []
    for f in _iter_output_files(d):
        rel_path = f.relative_to(d).as_posix()
        ext = f.suffix.lower()
        if ext in _INLINE_EXTS or ext == ".docx":
            try:
                if ext == ".docx":
                    content = _parse_docx(f)
                else:
                    content = f.read_text(encoding="utf-8")
                inline_parts.append(f"--- {rel_path} ---\n{content}")
            except Exception:
                continue
        else:
            data_paths.append(str(f.resolve()))
    return {"inline": "\n\n".join(inline_parts), "paths": data_paths}


def _sync_to_storage(storage: Storage, out: Path, meta_path: Path):
    """Upload JSONL + meta to remote storage."""
    try:
        storage.put(out.name, out)
        storage.put(meta_path.name, meta_path)
    except Exception as e:
        log.error(f"  [SYNC ERROR] {e}")


def _upload_snapshot_and_cleanup(
    storage: Storage,
    out_name: str,
    out_snapshot: Path,
    meta_name: str,
    meta_snapshot: Path,
):
    """Upload snapshot files and remove temporary copies."""
    try:
        storage.put(out_name, out_snapshot)
        storage.put(meta_name, meta_snapshot)
    except Exception as e:
        log.error(f"  [SYNC ERROR] {e}")
    finally:
        for path in (out_snapshot, meta_snapshot):
            try:
                path.unlink(missing_ok=True)
            except OSError:
                pass


def _schedule_sync_snapshot(storage: Storage, out: Path, meta_path: Path):
    """Snapshot current files before async upload to avoid read/write races."""
    from eval.storage import fire_and_forget

    token = uuid.uuid4().hex
    sync_dir = out.parent / ".sync_tmp"
    sync_dir.mkdir(parents=True, exist_ok=True)
    out_snapshot = sync_dir / f"{out.name}.{token}.snap"
    meta_snapshot = sync_dir / f"{meta_path.name}.{token}.snap"
    shutil.copy2(out, out_snapshot)
    shutil.copy2(meta_path, meta_snapshot)
    fire_and_forget(
        _upload_snapshot_and_cleanup,
        storage,
        out.name,
        out_snapshot,
        meta_path.name,
        meta_snapshot,
    )


def _run_task(
    llm_module,
    task: dict,
    refs: dict,
    llm_id: str,
    criterion_workers: int | None = None,
) -> dict:
    """Run a single task: generate answer, collect output, judge, score. Thread-safe."""
    tid = task["id"]
    task_repl = PythonREPL()

    # Output dir from dataset
    output_dir_name = task.get("output_dir")
    task_out_dir = _resolve_output_dir(output_dir_name) if output_dir_name else None
    if task_out_dir:
        _clear_output_dir(task_out_dir)

    # Inject REPL and output dir into config
    config = {**task.get("config", {}), "_repl": task_repl}
    if task_out_dir:
        config["_output_dir"] = str(task_out_dir)

    try:
        raw_answer = llm_module.generate(task["task"], refs, config)
    except Exception as e:
        log.error(f"  [{tid}] LLM ERROR: {e}")
        raw_answer = f"[ERROR: {e}]"

    # Unpack dict return (text + metadata) or plain string
    if isinstance(raw_answer, dict):
        answer = raw_answer.get("text", "")
        llm_metadata = raw_answer.get("metadata")
    else:
        answer = raw_answer
        llm_metadata = None

    artifact_bundle = build_judge_artifact_bundle(output_dir_name) if task_out_dir else {
        "artifact_root": None,
        "prompt_context": None,
        "repl_seed": {},
    }

    answer_len = len(answer)
    cost_info = f", ${llm_metadata['cost_usd']:.4f}" if llm_metadata and 'cost_usd' in llm_metadata else ""
    meta_info = f" ({llm_metadata['total_turns']} turns, {llm_metadata['usage']['total_tokens']} tokens{cost_info})" if llm_metadata else ""
    log.info(f"  [{tid}] answer_len={answer_len}{meta_info}")

    rubric = task["rubric"]
    n_m = len(rubric.get("mandatory", []))
    n_g = len(rubric.get("good_to_have", []))
    n_i = len(rubric.get("ideal", []))
    log.info(f"  [{tid}] [JUDGING] {n_m + n_g + n_i} criteria (mandatory={n_m}, good_to_have={n_g}, ideal={n_i})")
    eval_results = judge_rubric(
        task["task"],
        answer,
        task["rubric"],
        criterion_workers=criterion_workers,
        repl_seed=artifact_bundle["repl_seed"] or None,
        output_dir=artifact_bundle["artifact_root"],
        artifact_context=artifact_bundle["prompt_context"],
    )
    task_score = score_rubric(eval_results["mandatory"], eval_results["good_to_have"], eval_results["ideal"])

    m_pass = sum(eval_results["mandatory"])
    m_total = len(eval_results["mandatory"])
    log.info(f"  [{tid}] mandatory={m_pass}/{m_total} score={task_score:.4f}")

    return {
        "row": {
            **task,
            "llm_answer": answer,
            "eval": {**eval_results, "score": task_score},
            "llm_id": llm_id,
            "eval_timestamp": datetime.now(timezone.utc).isoformat(),
        },
        "score": task_score,
        "llm_metadata": llm_metadata,
        "m_pass": m_pass,
        "m_total": m_total,
    }


def _detect_concurrency() -> int:
    """Auto-detect thread count: cpu_count capped at 8, minimum 1."""
    return min(os.cpu_count() or 4, 8)


def _init_weave(project: str | None):
    if not project:
        return None
    try:
        import weave  # type: ignore
    except ImportError:
        log.warning("weave_project is set but weave is not installed; skipping Weave tracing")
        return None
    try:
        # weave.init expects "entity/project" or just "project" with WANDB_ENTITY set
        entity = os.environ.get("WANDB_ENTITY") or os.environ.get("WEAVE_ENTITY")
        if entity and "/" not in project:
            project = f"{entity}/{project}"
        return weave.init(project)
    except Exception as exc:
        log.warning(f"Failed to initialize Weave: {type(exc).__name__}: {exc}")
        return None


def run_eval(
    llm_module,
    dataset_path: Path | None = None,
    output_path: Path | None = None,
    task_ids: set[str] | None = None,
    storage: Storage | None = None,
    wandb_project: str | None = None,
    weave_project: str | None = None,
    concurrency: int | None = None,
) -> Path:
    """Run full eval with crash-resume, parallel task execution, S3 persistence, and W&B tracking."""
    tasks = load_dataset(dataset_path)
    if task_ids:
        tasks = [t for t in tasks if t["id"] in task_ids]
    llm_id = llm_module.LLM_ID
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    workers = concurrency or _detect_concurrency()
    judge_workers_env = os.getenv("JUDGE_CRITERION_WORKERS")
    if judge_workers_env:
        judge_workers = max(1, int(judge_workers_env))
    else:
        # Cap total nested judge threads to avoid excessive fan-out.
        judge_workers = max(1, min(16, 32 // max(1, workers)))

    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    out = output_path or (RESULTS_DIR / f"{llm_id}_{ts}.jsonl")
    meta_path = out.with_suffix(".meta.json")
    run_id = out.stem

    # Try restoring from remote storage if no local file
    if storage and not out.exists():
        if storage.get(out.name, out):
            log.info(f"Restored {out.name} from remote storage")
        storage.get(meta_path.name, meta_path)

    # Resume: skip already-completed tasks
    completed_ids = _load_completed_ids(out)
    all_task_count = len(tasks)
    remaining = [t for t in tasks if t["id"] not in completed_ids]

    if completed_ids:
        log.info(f"Resuming: {len(completed_ids)}/{all_task_count} already done, {len(remaining)} remaining")

    log.info(
        f"Eval: {llm_id} | {len(remaining)}/{all_task_count} tasks | "
        f"concurrency={workers}, judge_concurrency={judge_workers} | -> {out}"
    )

    # Weave init (auto-patches OpenAI/Anthropic clients)
    weave_client = _init_weave(weave_project or os.environ.get("WEAVE_PROJECT"))

    # W&B init
    wb_run = None
    if wandb_project:
        import wandb
        wb_run = wandb.init(
            project=wandb_project, name=run_id, id=run_id,
            resume="allow",
            config={"llm_id": llm_id, "total_tasks": all_task_count, "dataset": str(dataset_path)},
        )

    _write_meta(meta_path,
        run_id=run_id, llm_id=llm_id, status="running", output=str(out),
        total=all_task_count, completed=len(completed_ids),
        started_at=datetime.now(timezone.utc).isoformat())

    # Preload references for remaining tasks
    task_refs = [read_reference_files(t.get("reference_files")) for t in remaining]

    # Collect scores from already-completed tasks
    scores = []
    if out.exists():
        for line in open(out):
            try:
                scores.append(json.loads(line)["eval"]["score"])
            except (json.JSONDecodeError, KeyError):
                pass

    write_lock = threading.Lock()
    completed_count = len(completed_ids)

    def _on_result(result: dict):
        """Thread-safe callback: write result, update meta, log to W&B."""
        nonlocal completed_count
        with write_lock:
            completed_count += 1
            done = completed_count
            scores.append(result["score"])

            with open(out, "a") as f:
                f.write(json.dumps(result["row"]) + "\n")
                f.flush()

            tid = result["row"]["id"]
            _meta_update = {"completed": done, "current_task": tid}
            if result["llm_metadata"]:
                _meta_update[f"tasks.{tid}"] = result["llm_metadata"]
            _write_meta(meta_path, **_meta_update)

            if storage:
                _schedule_sync_snapshot(storage, out, meta_path)

            if wb_run:
                import wandb
                wb_data = {
                    "score": result["score"],
                    "mandatory_pass_rate": result["m_pass"] / result["m_total"] if result["m_total"] else 0,
                    "completed": done,
                }
                if result["llm_metadata"]:
                    wb_data["tokens"] = result["llm_metadata"].get("usage", {}).get("total_tokens", 0)
                    wb_data["turns"] = result["llm_metadata"].get("total_turns", 0)
                    if "cost_usd" in result["llm_metadata"]:
                        wb_data["cost_usd"] = result["llm_metadata"]["cost_usd"]
                wandb.log(wb_data, step=done)

    # Run tasks in parallel
    if len(remaining) <= 1 or workers <= 1:
        # Sequential for single task or concurrency=1
        for i, task in enumerate(remaining):
            log.info(f"[{len(completed_ids) + i + 1}/{all_task_count}] {task['id']}")
            result = _run_task(
                llm_module,
                task,
                task_refs[i],
                llm_id,
                criterion_workers=judge_workers,
            )
            _on_result(result)
    else:
        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {}
            for i, task in enumerate(remaining):
                log.info(f"[queued] {task['id']}")
                future = executor.submit(
                    _run_task,
                    llm_module,
                    task,
                    task_refs[i],
                    llm_id,
                    criterion_workers=judge_workers,
                )
                futures[future] = task["id"]

            for future in as_completed(futures):
                tid = futures[future]
                try:
                    result = future.result()
                    _on_result(result)
                except Exception as e:
                    log.error(f"[{tid}] TASK FAILED: {e}")

    avg = sum(scores) / len(scores) if scores else 0
    _write_meta(meta_path, status="completed", completed=all_task_count,
        avg_score=round(avg, 4), finished_at=datetime.now(timezone.utc).isoformat())

    # Final sync
    if storage:
        _sync_to_storage(storage, out, meta_path)

    # W&B finalize
    if wb_run:
        import wandb
        total_cost = 0
        total_tokens = 0
        if out.exists():
            for line in open(out):
                try:
                    r = json.loads(line)
                    meta = r.get("llm_metadata") or (r.get("eval", {}).get("metadata"))
                    if not meta:
                        continue
                    total_cost += meta.get("cost_usd", 0)
                    total_tokens += meta.get("usage", {}).get("total_tokens", 0)
                except (json.JSONDecodeError, KeyError):
                    pass
        wandb.summary["avg_score"] = avg
        wandb.summary["total_cost_usd"] = total_cost
        wandb.summary["total_tokens"] = total_tokens
        wandb.summary["total_tasks"] = all_task_count
        artifact = wandb.Artifact(f"results-{llm_id}", type="eval-results")
        artifact.add_file(str(out))
        wb_run.log_artifact(artifact)
        wandb.finish()

    if weave_client is not None:
        try:
            import weave  # type: ignore
            weave.finish()
        except Exception:
            pass

    log.info(f"Done. avg_score={avg:.4f} | {out}")
    return out
